from docx import Document
from docx.shared import Inches
from pylatex import Document as LaTeXDoc
from pylatex.utils import NoEscape
import os


def latex_to_image_pylatex(latex_str, filename):
    """使用pylatex生成更高质量的LaTeX公式图片"""
    # 创建临时LaTeX文档
    doc = LaTeXDoc()
    doc.preamble.append(NoEscape(r'\usepackage{amsmath}'))
    doc.preamble.append(NoEscape(r'\usepackage{amssymb}'))

    # 添加公式
    doc.append(NoEscape(f'${latex_str}$'))

    # 生成PDF
    doc.generate_pdf('temp_formula', clean_tex=False)

    # 这里需要将PDF转换为PNG
    # 可以使用pdf2image或其他工具
    # 这里简化处理，实际使用时需要添加转换代码


def create_word_with_pylatex():
    doc = Document()
    doc.add_heading('高质量LaTeX公式文档', 0)

    formulas = [
        r"E = mc^2",
        r"\int_{-\infty}^{\infty} e^{-x^2} dx = \sqrt{\pi}",
        r"\frac{\partial f}{\partial t} + \nabla \cdot \mathbf{J} = 0"
    ]

    for i, formula in enumerate(formulas):
        doc.add_paragraph(f'公式 {i + 1}: {formula}')
        # 在实际使用中，这里会插入由pylatex生成的高质量公式图片
        # img_filename = f'formula_{i}.png'
        # latex_to_image_pylatex(formula, img_filename)
        # doc.add_picture(img_filename, width=Inches(4))
        # os.remove(img_filename)

    doc.save('high_quality_formulas.docx')
    print("Word文档已生成: high_quality_formulas.docx")


if __name__ == "__main__":
    create_word_with_pylatex()